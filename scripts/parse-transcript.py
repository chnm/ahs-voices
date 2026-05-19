#!/usr/bin/env python3
"""Parse oral history transcript from docx table format into timestamped text.

Usage:
    python3 scripts/parse-transcript.py <input.docx> [--json]

Outputs the transcript in [HH:MM:SS] format suitable for the oralhistory:transcript
field in Omeka S, plus extracted metadata.
"""

import sys
import re
import json

try:
    import docx
except ImportError:
    print("Error: python-docx required. Install with: uv pip install python-docx")
    sys.exit(1)


def parse_transcript(path):
    doc = docx.Document(path)

    # Extract cover page metadata from paragraphs before the first table
    metadata = {
        "narrator": "",
        "interviewer": "",
        "location": "",
        "date": "",
        "transcriber": "",
        "summary": "",
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Interviewer:"):
            metadata["interviewer"] = text.replace("Interviewer:", "").strip()
        elif text.startswith("Narrator:"):
            metadata["narrator"] = text.replace("Narrator:", "").strip()
        elif text.startswith("Location of Interview:"):
            metadata["location"] = text.replace("Location of Interview:", "").strip()
        elif text.startswith("Date of Interview:"):
            metadata["date"] = text.replace("Date of Interview:", "").strip()
        elif text.startswith("Transcriber:"):
            metadata["transcriber"] = text.replace("Transcriber:", "").strip()

    # Extract summary (paragraphs between "Summary:" and the next section)
    in_summary = False
    summary_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "Summary:":
            in_summary = True
            continue
        if in_summary:
            if text.startswith("THE ") or text.startswith("OFFICE") or text.startswith("Oral History Summary"):
                break
            if text:
                summary_parts.append(text)
    metadata["summary"] = " ".join(summary_parts)

    # Extract keywords from table 1
    keywords = {}
    if len(doc.tables) > 1:
        kw_table = doc.tables[1]
        for row in kw_table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                keywords[cells[0]] = cells[1]
    metadata["keywords"] = keywords

    # Parse transcript from the last large table
    # Find the transcript table (2-column, with timestamps in left column)
    transcript_table = None
    for table in doc.tables:
        if len(table.columns) == 2:
            first_cell = table.rows[0].cells[0].text.strip()
            if re.search(r'\d{2}:\d{2}:\d{2}', first_cell):
                transcript_table = table
                break

    if not transcript_table:
        print("Error: Could not find transcript table", file=sys.stderr)
        return metadata, []

    # The transcript is often a single row with all content in two cells
    # Left: speaker names + timestamps, Right: spoken text
    # Both use paragraph breaks to separate turns

    left_cell = transcript_table.rows[0].cells[0]
    right_cell = transcript_table.rows[0].cells[1]

    # Parse left column into speaker entries
    left_paras = [p.text.strip() for p in left_cell.paragraphs]
    speakers = []
    i = 0
    while i < len(left_paras):
        text = left_paras[i]
        if text and not re.match(r'^\d{2}:\d{2}:\d{2}', text):
            # Speaker name found, look for timestamp
            name = text
            ts = ""
            for j in range(i + 1, min(i + 3, len(left_paras))):
                if re.match(r'^\d{2}:\d{2}:\d{2}', left_paras[j]):
                    ts = left_paras[j].strip()
                    # Count empty paragraphs after timestamp to know how many
                    # right-column paragraphs belong to this speaker
                    empty_count = 0
                    for k in range(j + 1, len(left_paras)):
                        if left_paras[k]:
                            break
                        empty_count += 1
                    speakers.append((name, ts, empty_count))
                    i = j
                    break
        i += 1

    # Parse right column paragraphs
    right_paras = [p.text.strip() for p in right_cell.paragraphs]

    # Build transcript entries by consuming right paragraphs per speaker
    entries = []
    right_idx = 0
    for name, ts, empty_after in speakers:
        # Collect paragraphs for this speaker turn
        # The number of right paragraphs roughly matches the empty lines in the left
        # But it's easier to just collect until we've consumed the right amount
        text_parts = []
        # Consume non-empty paragraphs, then skip empty ones
        while right_idx < len(right_paras):
            if right_paras[right_idx]:
                text_parts.append(right_paras[right_idx])
                right_idx += 1
            else:
                right_idx += 1  # skip empty
                break

        # Continue collecting if there are more non-empty paragraphs
        # that belong to this speaker (multi-paragraph responses)
        # We peek ahead: if the next speaker hasn't started yet,
        # keep consuming
        while right_idx < len(right_paras) and right_paras[right_idx]:
            text_parts.append(right_paras[right_idx])
            right_idx += 1
        # Skip trailing empties
        while right_idx < len(right_paras) and not right_paras[right_idx]:
            right_idx += 1

        entries.append({
            "speaker": name,
            "timestamp": ts,
            "text": "\n\n".join(text_parts),
        })

    return metadata, entries


def format_transcript(entries):
    """Format transcript entries as [HH:MM:SS] text for Omeka."""
    lines = []
    for entry in entries:
        ts = entry["timestamp"]
        speaker = entry["speaker"]
        text = entry["text"]
        lines.append(f"[{ts}] {speaker}")
        lines.append(text)
        lines.append("")
    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/parse-transcript.py <input.docx> [--json]")
        sys.exit(1)

    path = sys.argv[1]
    as_json = "--json" in sys.argv

    metadata, entries = parse_transcript(path)

    if as_json:
        output = {
            "metadata": metadata,
            "transcript": format_transcript(entries),
            "entries": entries,
        }
        print(json.dumps(output, indent=2, ensure_ascii=False))
    else:
        print("=== METADATA ===")
        for k, v in metadata.items():
            if k != "keywords":
                print(f"  {k}: {v}")
        if metadata.get("keywords"):
            print("  keywords:")
            for k, v in metadata["keywords"].items():
                print(f"    {k}: {v}")
        print()
        print("=== TRANSCRIPT ===")
        print(format_transcript(entries))
